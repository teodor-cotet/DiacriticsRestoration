import docx
import subprocess

if __name__ == "__main__":
    filee = 'test_restore/Informare Sisteme noi - sindicat și angajați_v1 (1).docx'
    tmp_file = "test_restore/tmp.txt"
    tmp_res = 'test_restore/tmp_res.txt'
    #print(get_docx_text(filee))
    doc = docx.Document(filee)
    #output = open("output_"+cur_date+"_.txt", "w")
    for i in range(len(doc.paragraphs)):
        #if 'sea' in paragraph.text:
        par = doc.paragraphs[i]
        with open(tmp_file, 'w') as f:
            f.write(par.text)
        subprocess.call(["python3", "model_diacritice.py", "-restore", tmp_file, "-load" ,"only_chars_win31_256_64_55","-classes" , "5"])
        txt_res = ""
        print('x')
        with open(tmp_res, 'r') as f:
            for line in f:
                txt_res += line
        par.text = txt_res
    doc.save('test_restore/modified.docx')
